"""MCP Word Server — python-docx based Word operations.

Tools:
  word_read    — Read text from a Word document
  word_write   — Write content to a Word document (overwrite)
  word_create  — Create a new Word document
  word_replace — Find-and-replace placeholders {{name}} → value
  word_merge   — Merge multiple documents into one
"""

import json
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from mcp_protocol import McpServer, McpTool, log_stderr
from repl_skin import ReplSkin

skin = ReplSkin("word", version="1.0.0")


class WordReadTool(McpTool):
    def __init__(self):
        super().__init__(
            name="word_read",
            description="Read text from a Word document",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to .docx file"},
                },
                "required": ["path"],
            },
        )

    def execute(self, arguments: dict) -> dict:
        from docx import Document
        path = arguments["path"]
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        tables = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)
        return {"text": full_text, "paragraphs": paragraphs, "tables": tables, "para_count": len(paragraphs)}


class WordWriteTool(McpTool):
    def __init__(self):
        super().__init__(
            name="word_write",
            description="Write content to a Word document (overwrite)",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string", "description": "Text content to write"},
                    "mode": {"type": "string", "default": "overwrite"},
                },
                "required": ["path", "content"],
            },
        )

    def execute(self, arguments: dict) -> dict:
        from docx import Document
        path = arguments["path"]
        content = arguments["content"]
        mode = arguments.get("mode", "overwrite")

        if mode == "overwrite" or not os.path.exists(path):
            doc = Document()
        else:
            doc = Document(path)

        # Clear existing paragraphs
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
        return {"path": path, "mode": mode, "lines": len(content.split("\n"))}


class WordCreateTool(McpTool):
    def __init__(self):
        super().__init__(
            name="word_create",
            description="Create a new Word document with title and content",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "title": {"type": "string"},
                    "content": {"type": "string"},
                },
                "required": ["path"],
            },
        )

    def execute(self, arguments: dict) -> dict:
        from docx import Document
        from docx.shared import Pt
        path = arguments["path"]
        title = arguments.get("title", "")
        content = arguments.get("content", "")
        doc = Document()
        if title:
            heading = doc.add_heading(title, level=1)
        if content:
            for line in content.split("\n"):
                doc.add_paragraph(line)
        doc.save(path)
        return {"path": path, "title": title}


class WordReplaceTool(McpTool):
    def __init__(self):
        super().__init__(
            name="word_replace",
            description="Find-and-replace placeholders in a Word document",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "find": {"type": "string", "description": "Text or {{placeholder}} to find"},
                    "replace": {"type": "string", "description": "Replacement text"},
                },
                "required": ["path", "find", "replace"],
            },
        )

    def execute(self, arguments: dict) -> dict:
        from docx import Document
        path = arguments["path"]
        find = arguments["find"]
        replace = arguments.get("replace", "")
        doc = Document(path)
        count = 0
        for p in doc.paragraphs:
            if find in p.text:
                # Replace in runs (handles formatting)
                for run in p.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
                # Also handle mixed formatting case
                if find in p.text:
                    full = p.text
                    replaced = full.replace(find, replace)
                    if replaced != full:
                        for run in p.runs:
                            run.text = ""
                        p.runs[0].text = replaced if p.runs else ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if find in run.text:
                                    run.text = run.text.replace(find, replace)
                                    count += 1
        doc.save(path)
        return {"path": path, "find": find, "replaced_count": count}


class WordMergeTool(McpTool):
    def __init__(self):
        super().__init__(
            name="word_merge",
            description="Merge multiple Word documents into one",
            input_schema={
                "type": "object",
                "properties": {
                    "files": {"type": "array", "description": "List of .docx file paths"},
                    "output": {"type": "string", "description": "Output file path"},
                },
                "required": ["files", "output"],
            },
        )

    def execute(self, arguments: dict) -> dict:
        from docx import Document
        from docx.enum.text import WD_BREAK
        files = arguments["files"]
        output = arguments["output"]
        merged = Document()
        for i, fpath in enumerate(files):
            if not os.path.exists(fpath):
                log_stderr(f"File not found: {fpath}")
                continue
            doc = Document(fpath)
            for p in doc.paragraphs:
                merged.add_paragraph(p.text)
            if i < len(files) - 1:
                merged.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        merged.save(output)
        return {"output": output, "files_merged": len(files)}


def main():
    server = McpServer("wf-word", "1.0.0")
    server.add_tool(WordReadTool())
    server.add_tool(WordWriteTool())
    server.add_tool(WordCreateTool())
    server.add_tool(WordReplaceTool())
    server.add_tool(WordMergeTool())
    server.run()


if __name__ == "__main__":
    main()
