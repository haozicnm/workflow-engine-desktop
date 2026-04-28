#!/usr/bin/env python3
"""
完整自动化管道（含 Playwright 浏览器自动化）
步骤：读Excel(A列) -> 网页查询 -> 写回Excel(B列) -> Word模板替换 -> 输出报告

用法：
  python run_full_pipeline.py [--headless] [--browser]
"""
import sys
import os
import datetime
import json
import argparse

sys.stdout.reconfigure(encoding='utf-8')

BASE = os.path.dirname(os.path.abspath(__file__))


def read_excel(path):
    """读取 Excel A 列数据"""
    from openpyxl import load_workbook
    wb = load_workbook(path)
    ws = wb.active
    data = []
    for row in ws.iter_rows(min_row=2, max_col=1, values_only=False):
        cell = row[0]
        if cell.value:
            data.append({"row": cell.row, "keyword": str(cell.value)})
    return data, wb, ws


def simulate_query(keywords):
    """模拟查询（不使用浏览器）"""
    results = {}
    categories = {
        "Python": ("编程语言", "广泛应用于AI、数据科学、Web开发"),
        "Rust": ("系统编程", "内存安全、零成本抽象、并发安全"),
        "Vue": ("前端框架", "渐进式JavaScript框架、响应式数据绑定"),
        "Tauri": ("桌面应用", "Rust后端+WebView前端、轻量级跨平台"),
        "自动化": ("自动化工具", "流程自动化、RPA、工作流引擎"),
    }
    for kw in keywords:
        keyword = kw["keyword"]
        row = kw["row"]
        category, desc = "未知", "未找到分类"
        for key, (cat, d) in categories.items():
            if key in keyword:
                category, desc = cat, d
                break
        results[row] = {
            "keyword": keyword,
            "result": f"[{category}] {desc}",
            "source": "模拟查询",
        }
        print(f"  {keyword} -> [{category}] {desc}")
    return results


def query_web(keywords, headless=True):
    """使用 Playwright 浏览器查询"""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("  [WARNING] playwright not installed, using simulation")
        return simulate_query(keywords)

    results = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        page = browser.new_page()
        for kw in keywords:
            try:
                print(f"  Querying: {kw['keyword']}...")
                page.goto("https://www.baidu.com", timeout=15000)
                page.wait_for_selector("#kw", timeout=5000)
                page.fill("#kw", kw["keyword"])
                page.click("#su")
                page.wait_for_timeout(2000)
                result_text = page.evaluate("""
                    () => {
                        const results = document.querySelectorAll('.result h3 a');
                        return results.length > 0 ? results[0].textContent.trim() : 'No result';
                    }
                """)
                results[kw["row"]] = {
                    "keyword": kw["keyword"],
                    "result": result_text[:100],
                    "source": "Baidu Search",
                }
                print(f"    -> {result_text[:60]}...")
            except Exception as e:
                print(f"    -> Query failed: {e}")
                results[kw["row"]] = {
                    "keyword": kw["keyword"],
                    "result": f"Failed: {str(e)[:50]}",
                    "source": "Error",
                }
        browser.close()
    return results


def write_results_to_excel(wb, ws, results, path):
    """写入结果到 Excel B 列"""
    for row_num, data in results.items():
        ws[f"B{row_num}"] = data["result"]
    wb.save(path)


def generate_report(results, template_path, output_path, date_str, total):
    """使用 Word 模板生成报告"""
    from docx import Document

    doc = Document(template_path)
    placeholders = {
        "{{DATE}}": date_str,
        "{{TOTAL}}": str(total),
    }

    for i, (row_num, data) in enumerate(sorted(results.items()), start=1):
        placeholders[f"{{{{ROW{i}_NUM}}}}"] = str(i)
        placeholders[f"{{{{ROW{i}_KEYWORD}}}}"] = data["keyword"]
        placeholders[f"{{{{ROW{i}_RESULT}}}}"] = data["result"]

    for para in doc.paragraphs:
        for old, new in placeholders.items():
            if old in para.text:
                for run in para.runs:
                    run.text = run.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in placeholders.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--headless", action="store_true")
    parser.add_argument("--browser", action="store_true")
    parser.add_argument("--excel", default=f"{BASE}/test_data.xlsx")
    parser.add_argument("--template", default=f"{BASE}/report_template.docx")
    parser.add_argument("--output", default=f"{BASE}/report_output.docx")
    args = parser.parse_args()

    print("=" * 60)
    print("  Automation Pipeline: Excel -> Web -> Excel -> Word")
    print("=" * 60)

    # 1. Read Excel
    print("\n[1/4] Reading Excel column A...")
    data, wb, ws = read_excel(args.excel)
    for d in data:
        print(f"  A{d['row']}: {d['keyword']}")
    print(f"  Total: {len(data)} items")

    # 2. Query
    print("\n[2/4] Processing...")
    if args.browser:
        results = query_web(data, headless=args.headless)
    else:
        results = simulate_query(data)

    # 3. Write back
    print("\n[3/4] Writing results to Excel column B...")
    write_results_to_excel(wb, ws, results, args.excel)
    print(f"  Written {len(results)} results")

    # 4. Generate Word report
    print("\n[4/4] Generating Word report...")
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    generate_report(results, args.template, args.output, date_str, len(data))
    print(f"  Report: {args.output}")

    print("\n" + "=" * 60)
    print("  Pipeline completed!")
    print(f"  Excel: {args.excel}")
    print(f"  Report: {args.output}")
    print("=" * 60)

    output = {
        "success": True,
        "total": len(data),
        "excel_path": args.excel,
        "report_path": args.output,
    }
    print(f"\n__RESULT_JSON__:{json.dumps(output, ensure_ascii=False)}")
    return 0


if __name__ == "__main__":
    sys.exit(main() or 0)
