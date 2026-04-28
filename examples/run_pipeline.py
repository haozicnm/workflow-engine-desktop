#!/usr/bin/env python3
"""
完整自动化管道演示
读取Excel(A列) -> 处理 -> 写回Excel(B列) -> Word模板替换 -> 输出报告
"""
import sys
import os
import datetime

sys.stdout.reconfigure(encoding='utf-8')

from openpyxl import load_workbook
from docx import Document

BASE = "C:/Users/haozi/.openclaw/workspace/workflow-engine-desktop/examples"
EXCEL_PATH = f"{BASE}/test_data.xlsx"
TEMPLATE_PATH = f"{BASE}/report_template.docx"
REPORT_PATH = f"{BASE}/report_output.docx"

def process_keyword(keyword):
    """模拟网页查询：实际可以调用搜索引擎或API"""
    # 这里简单模拟：返回关键词的长度和特征
    length = len(keyword)
    if "Python" in keyword:
        category = "编程语言"
    elif "Rust" in keyword:
        category = "系统编程"
    elif "Vue" in keyword:
        category = "前端框架"
    elif "Tauri" in keyword:
        category = "桌面应用"
    elif "自动化" in keyword:
        category = "自动化工具"
    else:
        category = "其他"

    return f"{category} | {length}字符"

def main():
    print("=" * 50)
    print("工作流自动化管道 - 端到端演示")
    print("=" * 50)

    # ===== 步骤1: 读取Excel =====
    print("\n[步骤1] 读取 Excel A列数据...")
    wb = load_workbook(EXCEL_PATH)
    ws = wb.active
    data = []
    for row in ws.iter_rows(min_row=2, max_col=1, values_only=False):
        cell = row[0]
        if cell.value:
            data.append((cell.row, cell.value))
            print(f"  A{cell.row}: {cell.value}")
    print(f"  共读取 {len(data)} 条数据")

    # ===== 步骤2: 处理数据（模拟网页查询） =====
    print("\n[步骤2] 处理每条数据...")
    results = {}
    for row_num, keyword in data:
        result = process_keyword(keyword)
        results[row_num] = result
        print(f"  {keyword} -> {result}")

    # ===== 步骤3: 写回Excel B列 =====
    print("\n[步骤3] 写入结果到 B列...")
    for row_num, result in results.items():
        ws[f"B{row_num}"] = result
    wb.save(EXCEL_PATH)
    print(f"  已保存到 {EXCEL_PATH}")

    # ===== 步骤4: 生成Word报告 =====
    print("\n[步骤4] 生成 Word 报告...")
    doc = Document(TEMPLATE_PATH)

    today = datetime.date.today().strftime("%Y-%m-%d")
    placeholders = {
        "{{DATE}}": today,
        "{{TOTAL}}": str(len(data)),
    }

    for i, (row_num, keyword) in enumerate(data, start=1):
        placeholders[f"{{{{ROW{i}_NUM}}}}"] = str(i)
        placeholders[f"{{{{ROW{i}_KEYWORD}}}}"] = keyword
        placeholders[f"{{{{ROW{i}_RESULT}}}}"] = results[row_num]

    # 替换所有段落中的占位符
    for para in doc.paragraphs:
        for old, new in placeholders.items():
            if old in para.text:
                for run in para.runs:
                    run.text = run.text.replace(old, new)

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in placeholders.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)

    doc.save(REPORT_PATH)
    print(f"  报告已生成: {REPORT_PATH}")

    # ===== 完成 =====
    print("\n" + "=" * 50)
    print("管道执行完成！")
    print(f"  Excel: {EXCEL_PATH}")
    print(f"  报告:  {REPORT_PATH}")
    print("=" * 50)

    return 0

if __name__ == "__main__":
    sys.exit(main() or 0)
