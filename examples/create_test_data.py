#!/usr/bin/env python3
"""
创建测试 Excel 和 Word 模板文件
用于验证工作流自动化管道
"""
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')
from openpyxl import Workbook
from docx import Document

os.makedirs("C:/Users/haozi/.openclaw/workspace/workflow-engine-desktop/examples", exist_ok=True)
EXCEL_PATH = "C:/Users/haozi/.openclaw/workspace/workflow-engine-desktop/examples/test_data.xlsx"
TEMPLATE_PATH = "C:/Users/haozi/.openclaw/workspace/workflow-engine-desktop/examples/report_template.docx"
REPORT_PATH = "C:/Users/haozi/.openclaw/workspace/workflow-engine-desktop/examples/report_output.docx"

# ─── 创建测试 Excel ───

wb = Workbook()
ws = wb.active
ws.title = "数据"
ws["A1"] = "查询关键词"
ws["B1"] = "查询结果"

test_data = [
    "Python编程",
    "Rust语言入门",
    "Vue.js框架",
    "Tauri桌面应用",
    "工作流自动化",
]

for i, item in enumerate(test_data, start=2):
    ws[f"A{i}"] = item
    ws[f"B{i}"] = ""  # 待填充

wb.save(EXCEL_PATH)
print(f"✅ Excel 已创建: {EXCEL_PATH}")
print(f"   共 {len(test_data)} 条数据")

# ─── 创建 Word 模板 ───

doc = Document()
doc.add_heading("自动化查询报告", level=1)

doc.add_paragraph("报告生成日期：{{DATE}}")
doc.add_paragraph("")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "序号"
hdr[1].text = "查询关键词"
hdr[2].text = "查询结果"

for i in range(len(test_data)):
    row = table.add_row().cells
    row[0].text = f"{{{{ROW{i+1}_NUM}}}}"
    row[1].text = f"{{{{ROW{i+1}_KEYWORD}}}}"
    row[2].text = f"{{{{ROW{i+1}_RESULT}}}}"

doc.add_paragraph("")
doc.add_paragraph("总结：共查询 {{TOTAL}} 条数据。")

doc.save(TEMPLATE_PATH)
print(f"✅ Word 模板已创建: {TEMPLATE_PATH}")

# ─── 验证读取 ───

from openpyxl import load_workbook

wb2 = load_workbook(EXCEL_PATH)
ws2 = wb2.active
print(f"\n📖 读取测试:")
for row in ws2.iter_rows(min_row=2, max_col=1, values_only=True):
    if row[0]:
        print(f"   A{ws2[row[0]].row}: {row[0]}")

print("\n管道步骤 1 (读Excel) ✅")
